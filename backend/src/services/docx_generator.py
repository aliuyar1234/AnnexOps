"""DOCX generator service for Annex IV documentation."""
from datetime import UTC, datetime
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Human-readable section titles
SECTION_TITLES = {
    "ANNEX4.GENERAL": "General Information",
    "ANNEX4.INTENDED_PURPOSE": "Intended Purpose",
    "ANNEX4.SYSTEM_DESCRIPTION": "System Description",
    "ANNEX4.RISK_MANAGEMENT": "Risk Management System",
    "ANNEX4.DATA_GOVERNANCE": "Data Governance",
    "ANNEX4.MODEL_TECHNICAL": "Model & Technical Documentation",
    "ANNEX4.PERFORMANCE": "Performance Metrics",
    "ANNEX4.HUMAN_OVERSIGHT": "Human Oversight",
    "ANNEX4.LOGGING": "Logging & Traceability",
    "ANNEX4.ACCURACY_ROBUSTNESS_CYBERSEC": "Accuracy, Robustness & Cybersecurity",
    "ANNEX4.POST_MARKET_MONITORING": "Post-Market Monitoring",
    "ANNEX4.CHANGE_MANAGEMENT": "Change Management",
}


def _format_field_name(field_name: str) -> str:
    """Convert snake_case field name to human-readable title.

    Args:
        field_name: Field name in snake_case

    Returns:
        Human-readable field title
    """
    # Replace underscores with spaces and capitalize each word
    return field_name.replace("_", " ").title()


def _format_field_value(value: Any) -> str:
    """Format field value for display in document.

    Args:
        value: Field value (can be string, list, dict, etc.)

    Returns:
        Formatted string representation
    """
    if value is None:
        return "Not specified"

    if isinstance(value, list):
        if not value:
            return "Not specified"
        # Join list items with newlines
        return "\n".join(f"• {item}" for item in value)

    if isinstance(value, dict):
        # Format dict as key-value pairs
        items = []
        for k, v in value.items():
            items.append(f"{_format_field_name(k)}: {v}")
        return "\n".join(items)

    return str(value)


def generate_annex_iv_document(
    system_info: dict[str, Any],
    version_info: dict[str, Any],
    sections: list[dict[str, Any]],
    evidence_items: list[dict[str, Any]],
) -> BytesIO:
    """Generate Annex IV documentation as DOCX file.

    Args:
        system_info: System information dict with keys: id, name, hr_use_case_type, intended_purpose
        version_info: Version information dict with keys: id, label, status, release_date
        sections: List of section dicts with keys: section_key, content, evidence_refs
        evidence_items: List of evidence item dicts with keys: id, title, type

    Returns:
        BytesIO buffer containing the generated DOCX file
    """
    doc = Document()

    # Add title page
    title = doc.add_heading("EU AI Act - Annex IV Technical Documentation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacer

    # System information
    doc.add_heading("System Information", level=1)
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid Accent 1'

    info_table.rows[0].cells[0].text = "System Name"
    info_table.rows[0].cells[1].text = system_info.get("name", "")

    info_table.rows[1].cells[0].text = "Version"
    info_table.rows[1].cells[1].text = version_info.get("label", "")

    info_table.rows[2].cells[0].text = "HR Use Case Type"
    info_table.rows[2].cells[1].text = system_info.get("hr_use_case_type", "")

    info_table.rows[3].cells[0].text = "Generated Date"
    info_table.rows[3].cells[1].text = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")

    doc.add_paragraph()  # Spacer

    # Organization information
    doc.add_heading("Organization Information", level=1)
    org_para = doc.add_paragraph()
    org_para.add_run("Organization: ").bold = True
    org_para.add_run(system_info.get("org_name", "Not specified"))

    doc.add_page_break()

    # Sort sections by section_key for deterministic ordering
    sorted_sections = sorted(sections, key=lambda s: s["section_key"])

    # Add each section
    for section in sorted_sections:
        section_key = section["section_key"]
        content = section.get("content", {})
        evidence_refs = section.get("evidence_refs", [])

        # Section heading
        section_title = SECTION_TITLES.get(section_key, section_key)
        doc.add_heading(section_title, level=1)

        # Add content fields as table
        if content:
            # Create table with field names and values
            field_table = doc.add_table(rows=len(content), cols=2)
            field_table.style = 'Light List Accent 1'

            # Sort content keys for deterministic ordering
            for idx, (field_name, field_value) in enumerate(sorted(content.items())):
                field_table.rows[idx].cells[0].text = _format_field_name(field_name)
                field_table.rows[idx].cells[1].text = _format_field_value(field_value)
        else:
            doc.add_paragraph("No content provided for this section.", style='Italic')

        doc.add_paragraph()  # Spacer

        # Add evidence references if any
        if evidence_refs:
            doc.add_heading("Evidence References", level=2)

            # Create evidence table
            ev_table = doc.add_table(rows=len(evidence_refs) + 1, cols=3)
            ev_table.style = 'Light Grid Accent 1'

            # Header row
            ev_table.rows[0].cells[0].text = "ID"
            ev_table.rows[0].cells[1].text = "Title"
            ev_table.rows[0].cells[2].text = "Type"

            # Find evidence items by ID (convert UUID to str for comparison)
            evidence_map = {str(ev["id"]): ev for ev in evidence_items}

            # Add each evidence reference
            for idx, ev_id in enumerate(sorted(evidence_refs, key=str), start=1):
                ev_id_str = str(ev_id)
                evidence = evidence_map.get(ev_id_str, {})

                ev_table.rows[idx].cells[0].text = ev_id_str[:8] + "..."  # Shortened ID
                ev_table.rows[idx].cells[1].text = evidence.get("title", "Unknown")
                ev_table.rows[idx].cells[2].text = evidence.get("type", "Unknown")

        doc.add_page_break()

    # Add evidence index appendix
    if evidence_items:
        doc.add_heading("Appendix: Evidence Index", level=1)

        # Create full evidence index table
        idx_table = doc.add_table(rows=len(evidence_items) + 1, cols=3)
        idx_table.style = 'Light Grid Accent 1'

        # Header row
        idx_table.rows[0].cells[0].text = "Evidence ID"
        idx_table.rows[0].cells[1].text = "Title"
        idx_table.rows[0].cells[2].text = "Type"

        # Sort evidence items by ID for deterministic ordering
        sorted_evidence = sorted(evidence_items, key=lambda e: str(e["id"]))

        # Add each evidence item
        for idx, evidence in enumerate(sorted_evidence, start=1):
            idx_table.rows[idx].cells[0].text = str(evidence["id"])
            idx_table.rows[idx].cells[1].text = evidence.get("title", "")
            idx_table.rows[idx].cells[2].text = evidence.get("type", "")

    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
